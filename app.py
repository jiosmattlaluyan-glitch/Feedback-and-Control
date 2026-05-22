import streamlit as st
import wikipedia
from docx import Document
from openai import OpenAI
import os

# =====================================
# 1. TOOL DEFINITIONS
# =====================================
class WikiSearchTool:
    def search(self, query):
        try:
            return wikipedia.summary(query, sentences=5)
        except Exception:
            return "Could not retrieve information from Wikipedia."

class DocumentTool:
    def create_document(self, title, content):
        try:
            doc = Document()
            doc.add_heading(title, level=1)
            doc.add_paragraph(content)

            # Ensure safe filename
            filename = f"{title.replace(' ', '_')}.docx"
            doc.save(filename)
            return filename
        except Exception as e:
            return f"Document Error: {e}"

# Initialize tool instances
wiki_tool = WikiSearchTool()
doc_tool = DocumentTool()

# =====================================
# 2. STREAMLIT INITIALIZATION & MEMORY
# =====================================
st.set_page_config(page_title="Agentic AI Chatbot", layout="centered")
st.title("🤖 Agentic AI Chatbot")
st.caption("Feedback and Control System — Powered by Local LLM")

# Initialize Session State Variables for chat history and stats
if "messages" not in st.session_state:
    st.session_state.messages = [
        {"role": "system", "content": """You are a helpful AI assistant.
Rules:
1. If the user asks to search, define, explain, or look up a topic, reply ONLY in this format: WIKI: topic
2. If the user asks to create, save, export, or generate a document, reply ONLY in this format: DOC: title | content
3. Otherwise answer normally.""" }
    ]

if "stats" not in st.session_state:
    st.session_state.stats = {
        "User Turns": 0,
        "Wiki Searches": 0,
        "Document Saves": 0,
        "AI Chats": 0
    }

# Sidebar for Tracking Statistics
st.sidebar.header("📊 Conversation Statistics")
for stat_name, val in st.session_state.stats.items():
    st.sidebar.metric(label=stat_name, value=val)

if st.sidebar.button("Reset Session & Clear Cache"):
    st.session_state.messages = [st.session_state.messages[0]]
    st.session_state.stats = {"User Turns": 0, "Wiki Searches": 0, "Document Saves": 0, "AI Chats": 0}
    st.rerun()

# =====================================
# 3. CLIENT SETUP (LM STUDIO)
# =====================================
client = OpenAI(base_url="http://127.0.0.1:1234/v1", api_key="lm-studio")

# =====================================
# 4. RENDER UI CHAT HISTORY
# =====================================
# Display prior messages (skip system prompt)
for msg in st.session_state.messages:
    if msg["role"] != "system":
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

# =====================================
# 5. CHAT LOGIC INPUT PROCESSING
# =====================================
if user_input := st.chat_input("Ask me something or request to save information..."):

    # 1. User Message Display
    st.session_state.stats["User Turns"] += 1
    with st.chat_message("user"):
        st.markdown(user_input)
    st.session_state.messages.append({"role": "user", "content": user_input})

    # 2. Call LLM for Route/Decision Handling
    try:
        response = client.chat.completions.create(
            model="meta-llama-3.1-8b-instruct",
            messages=st.session_state.messages,
            temperature=0.3
        )
        decision = response.choices[0].message.content.strip()

        # 3. Route Executions
        with st.chat_message("assistant"):

            # --- WIKIPEDIA ROUTE ---
            if decision.upper().startswith("WIKI:"):
                st.session_state.stats["Wiki Searches"] += 1
                topic = decision.replace("WIKI:", "").strip()

                st.write(f"🔍 *Intent Routed: Wikipedia Search for '{topic}'...*")
                wiki_result = wiki_tool.search(topic)

                st.markdown(wiki_result)
                st.session_state.messages.append({"role": "assistant", "content": wiki_result})

            # --- DOCUMENT ROUTE ---
            elif decision.upper().startswith("DOC:"):
                st.session_state.stats["Document Saves"] += 1

                try:
                    raw_data = decision.replace("DOC:", "").strip()
                    title, content = raw_data.split("|", 1)
                    title, content = title.strip(), content.strip()
                except ValueError:
                    title = "Generated_Document"
                    content = decision

                st.write(f"💾 *Intent Routed: Generating Document...*")
                filename = doc_tool.create_document(title, content)
                assistant_reply = f"✅ Document saved successfully as `{filename}`"

                st.markdown(assistant_reply)

                # Provide a direct download button inside the Streamlit Web Application interface!
                if os.path.exists(filename):
                    with open(filename, "rb") as file:
                        st.download_button(
                            label="⬇️ Download Word Document",
                            data=file,
                            file_name=filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

                st.session_state.messages.append({"role": "assistant", "content": assistant_reply})

            # --- NORMAL CONVERSATION ROUTE ---
            else:
                st.session_state.stats["AI Chats"] += 1

                # Create streaming response placeholder
                reply_placeholder = st.empty()
                full_reply = ""

                stream = client.chat.completions.create(
                    model="meta-llama-3.1-8b-instruct",
                    messages=st.session_state.messages,
                    temperature=0.7,
                    stream=True
                )

                for chunk in stream:
                    content = chunk.choices[0].delta.content
                    if content:
                        full_reply += content
                        reply_placeholder.markdown(full_reply)

                st.session_state.messages.append({"role": "assistant", "content": full_reply})

        # Force refresh metrics sidebar UI
        st.rerun()

    except Exception as e:
        st.error(f"Error connecting to local LLM Server: {e}. Ensure LM Studio is actively running on port 1234!")
